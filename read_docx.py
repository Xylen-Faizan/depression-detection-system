import docx
import sys
import codecs

def read_docx(filename):
    doc = docx.Document(filename)
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_data))
            
    with codecs.open('docx_extracted.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(full_text))

if __name__ == '__main__':
    read_docx('directory_structure_depression.docx')
